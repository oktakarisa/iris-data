# problem4_dataset_research.py
#
# Problem 4: Examining the Dataset Itself
#
# Goal:
#   Save background information about the Iris dataset
#   into a document (Word, PDF, and Markdown) inside the data/ folder.

import os
from docx import Document          # for Word output
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet  # for PDF

def get_dataset_info():
    """
    Background information about the Iris dataset.
    """
    return """
Iris Dataset Background
-----------------------
- The Iris dataset is one of the most famous datasets in statistics and machine learning.
- It was introduced in 1936 by Ronald A. Fisher, a British statistician.
- The dataset was originally used in his paper:
    "The use of multiple measurements in taxonomic problems".
- Fisher collected data about three species of Iris flowers:
    * Iris setosa
    * Iris versicolor
    * Iris virginica
- Each flower sample has 4 measurements (features):
    1. Sepal length
    2. Sepal width
    3. Petal length
    4. Petal width
- There are 150 samples total (50 per species).
- The dataset was originally collected from the Gaspé Peninsula in Canada.
- Why it matters:
    * It’s small and clean — easy to practice on.
    * It’s historically important — used in early examples of discriminant analysis.
    * It’s widely available in scikit-learn and other libraries.
"""

def save_as_word(text, path):
    doc = Document()
    doc.add_heading("Iris Dataset Background", level=1)
    for line in text.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(path)

def save_as_pdf(text, path):
    doc = SimpleDocTemplate(path)
    styles = getSampleStyleSheet()
    story = [Paragraph(line, styles["Normal"]) for line in text.splitlines() if line.strip()]
    doc.build(story)

def save_as_md(text, path):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)

if __name__ == "__main__":
    os.makedirs("data", exist_ok=True)
    text = get_dataset_info()

    save_as_word(text, "data/iris_dataset_info.docx")
    save_as_pdf(text, "data/iris_dataset_info.pdf")
    save_as_md(text, "data/iris_dataset_info.md")

    print("Dataset info saved in data/ folder (Word, PDF, Markdown).")
